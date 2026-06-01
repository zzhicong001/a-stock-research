#!/usr/bin/env python3
"""
将 bazi_final.docx 转换为 EPUB 格式，针对 iPhone 图书 App 阅读优化。
"""

import os, sys, shutil, zipfile
from pathlib import Path
from lxml import etree
import docx
from ebooklib import epub

# ── 配置 ──
DOCX_PATH = r"C:/Users/Administrator/Desktop/bazi_final.docx"
OUTPUT_EPUB = r"H:/buddy/华山八字教学总集.epub"
WORK_DIR = Path(r"H:/buddy/epub_build")
TITLE = "华山八字教学总集"
AUTHOR = "老陈说易"
LANGUAGE = "zh"

# ── 命名空间 ──
NS_RELS = 'http://schemas.openxmlformats.org/package/2006/relationships'
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ── CSS（iPhone 图书 App 优化） ──
CSS = """
@namespace epub "http://www.idpf.org/2007/ops";
body {
    font-family: -apple-system, "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
    font-size: 1.05em;
    line-height: 1.85;
    color: #1a1a1a;
    margin: 0;
    padding: 0.6em 0.8em;
    text-align: justify;
}
h1 {
    font-size: 1.6em; font-weight: 700; color: #8B0000; text-align: center;
    margin: 1.2em 0 0.6em 0; padding-bottom: 0.3em;
    border-bottom: 2px solid #8B0000; line-height: 1.3;
}
h2 {
    font-size: 1.3em; font-weight: 600; color: #333;
    margin: 1em 0 0.4em 0; padding: 0.2em 0 0.2em 0.5em;
    border-left: 4px solid #8B0000; line-height: 1.35;
}
h3 { font-size: 1.15em; font-weight: 600; color: #444; margin: 0.8em 0 0.3em 0; }
h4 { font-size: 1.05em; font-weight: 600; color: #555; margin: 0.7em 0 0.3em 0; }
p { margin: 0.4em 0; text-indent: 1.6em; }
p.no-indent { text-indent: 0; }
p.separator { text-align: center; color: #999; margin: 1em 0; text-indent: 0; font-size: 0.9em; }
p.source-link { font-size: 0.8em; color: #888; text-indent: 0; word-break: break-all; }
p.quoted { border-left: 3px solid #ccc; padding-left: 1em; color: #555; margin: 0.6em 0; }
img {
    max-width: 100%; height: auto; display: block;
    margin: 0.6em auto; border-radius: 4px;
}
.figure-caption { text-align: center; font-size: 0.85em; color: #777; text-indent: 0; }
em { font-style: italic; color: #555; }
strong { color: #222; }

@media (prefers-color-scheme: dark) {
    body { color: #ddd; }
    h1 { color: #ff6b6b; border-bottom-color: #ff6b6b; }
    h2 { color: #ccc; border-left-color: #ff6b6b; }
    h3, h4 { color: #bbb; }
    .source-link { color: #999; }
    em { color: #aaa; }
    strong { color: #eee; }
    .separator { color: #666; }
    .quoted { border-left-color: #555; color: #aaa; }
    .figure-caption { color: #888; }
}
"""

# ── 图片提取 ──

def get_image_map(docx_path):
    """从 docx 提取图片，返回 {rId: (filename, bytes, mime)}"""
    image_map = {}
    with zipfile.ZipFile(docx_path) as z:
        # 只读 document.xml.rels（所有图片引用在这里）
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in z.namelist():
            return image_map

        rels_xml = etree.fromstring(z.read(rels_path))
        # .rels 文件使用默认命名空间，直接用 Clark notation
        rels_ns = '{' + NS_RELS + '}'

        for rel in rels_xml.iter(rels_ns + 'Relationship'):
            rtype = rel.get('Type', '')
            if 'image' not in rtype.lower():
                continue
            rid = rel.get('Id')
            target = rel.get('Target')
            full_target = 'word/' + target

            if full_target in z.namelist():
                ext = os.path.splitext(target)[1].lower()
                mime_map = {'.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                            '.png': 'image/png', '.gif': 'image/gif',
                            '.bmp': 'image/bmp', '.webp': 'image/webp'}
                mime = mime_map.get(ext, 'image/jpeg')
                image_map[rid] = (os.path.basename(target), z.read(full_target), mime)

    return image_map

# ── 判断函数 ──

def is_separator(text):
    if not text:
        return False
    stripped = text.replace('—', '').strip()
    return stripped == '' and len(text) > 10

def is_source_link(text):
    return text.startswith('原文链接：') or text.startswith('原文链接:')

def get_blips(xml_element):
    """查找所有 blip 元素"""
    return xml_element.findall('.//{%s}blip' % NS_A)

def get_inline_drawings(xml_element):
    """查找所有内联绘图"""
    return xml_element.findall('.//{%s}inline' % NS_WP)

def has_image_content(para_element):
    """判断段落是否包含图片"""
    return bool(get_blips(para_element) or get_inline_drawings(para_element))

# ── 主转换 ──

def convert():
    print("📖 读取 docx...")
    doc = docx.Document(DOCX_PATH)
    print(f"   段落数: {len(doc.paragraphs)}")

    print("🖼️  提取图片引用...")
    image_map = get_image_map(DOCX_PATH)
    print(f"   图片引用数: {len(image_map)}")

    # 清理工作目录
    if WORK_DIR.exists():
        shutil.rmtree(WORK_DIR)
    WORK_DIR.mkdir(parents=True)

    # ── 第一遍：分析结构 ──
    print("📊 分析文档结构...")

    chapter_boundaries = []
    current_start = None
    current_level = None
    current_title = None
    h1_count = 0
    h2_count = 0

    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ''
        text = para.text.strip() if para.text else ''
        style_lower = style.lower()

        # 跳过纯空段落（无文本、无图片）
        if not text and not has_image_content(para._element):
            continue

        if 'heading 1' in style_lower:
            if current_start is not None:
                chapter_boundaries.append((current_start, i, current_level, current_title))
            current_start, current_level, current_title = i, 1, text
            h1_count += 1
        elif 'heading 2' in style_lower:
            if current_start is not None:
                chapter_boundaries.append((current_start, i, current_level, current_title))
            current_start, current_level, current_title = i, 2, text
            h2_count += 1
        elif 'heading 3' in style_lower:
            if current_start is not None:
                chapter_boundaries.append((current_start, i, current_level, current_title))
            current_start, current_level, current_title = i, 3, text

    if current_start is not None:
        chapter_boundaries.append((current_start, len(doc.paragraphs), current_level, current_title))

    # 标题前的目录内容作为首章
    if chapter_boundaries and chapter_boundaries[0][0] > 0:
        chapter_boundaries.insert(0, (0, chapter_boundaries[0][0], 0, "目录"))

    print(f"   H1={h1_count}, H2={h2_count}, 章节={len(chapter_boundaries)}")

    # ── 第二遍：生成 EPUB ──
    print("📝 生成 EPUB...")

    book = epub.EpubBook()
    book.set_identifier('huashan-bazi-teaching-v1')
    book.set_title(TITLE)
    book.set_language(LANGUAGE)
    book.add_author(AUTHOR)
    book.add_metadata('DC', 'description', '华山八字教学全系列，共557篇文章，老陈说易编著')
    book.add_metadata('DC', 'date', '2026-05-31')

    # CSS
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css",
                            media_type="text/css", content=CSS.encode('utf-8'))
    book.add_item(nav_css)

    epub_chapters = []
    spine = ['nav']
    toc_entries = []
    used_images = set()
    img_counter = 0

    for ch_idx, (start, end, level, title) in enumerate(chapter_boundaries):
        safe_id = f"ch{ch_idx:04d}"

        h_tag = max(level, 1) if level > 0 else 1
        html_parts = [f'<h{h_tag}>{title}</h{h_tag}>\n']
        is_first = True

        for pi in range(start, end):
            if pi >= len(doc.paragraphs):
                break
            para = doc.paragraphs[pi]
            text = para.text.strip() if para.text else ''
            style = para.style.name if para.style else ''
            style_lower = style.lower()
            xml_elem = para._element

            # 跳过当前章节的标题段落（已单独输出）
            if pi == start and any(h in style_lower for h in ('heading 1', 'heading 2', 'heading 3')):
                continue

            # 跳过纯空段落
            if not text and not has_image_content(xml_elem):
                continue

            # ── 处理图片 ──
            blips = get_blips(xml_elem)
            if blips:
                for blip in blips:
                    embed = blip.get('{%s}embed' % NS_R)
                    if embed and embed in image_map:
                        fname, img_data, mime = image_map[embed]
                        if fname not in used_images:
                            epub_img = epub.EpubImage()
                            epub_img.file_name = f'images/{fname}'
                            epub_img.media_type = mime
                            epub_img.content = img_data
                            book.add_item(epub_img)
                            used_images.add(fname)
                            img_counter += 1
                        html_parts.append(f'<div class="figure"><img src="images/{fname}" alt="图片"/></div>\n')

            # ── 处理文本 ──
            if text:
                if is_separator(text):
                    html_parts.append('<p class="separator">———</p>\n')
                elif is_source_link(text):
                    html_parts.append(f'<p class="source-link">{text}</p>\n')
                elif 'heading 2' in style_lower:
                    html_parts.append(f'<h2>{text}</h2>\n')
                elif 'heading 3' in style_lower:
                    html_parts.append(f'<h3>{text}</h3>\n')
                elif 'heading 4' in style_lower:
                    html_parts.append(f'<h4>{text}</h4>\n')
                else:
                    cls = 'no-indent' if is_first else ''
                    html_parts.append(f'<p class="{cls}">{text}</p>\n')
                    is_first = False

        body_content = "\n".join(html_parts)
        chapter = epub.EpubHtml(title=title, file_name=f'text/{safe_id}.xhtml', lang=LANGUAGE)
        chapter.content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="{LANGUAGE}">
<head><meta charset="UTF-8"/><title>{title}</title>
<link rel="stylesheet" type="text/css" href="../style/nav.css"/></head>
<body>
{body_content}
</body></html>'''.encode('utf-8')

        book.add_item(chapter)
        epub_chapters.append(chapter)
        spine.append(chapter)
        toc_entries.append(epub.Link(f'text/{safe_id}.xhtml', title, safe_id))

    # ── TOC & Spine ──
    book.toc = toc_entries
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # ── 写出 ──
    print(f"💾 写入 EPUB ({len(epub_chapters)} 章节, {img_counter} 图片)...")
    epub.write_epub(OUTPUT_EPUB, book)

    file_size = os.path.getsize(OUTPUT_EPUB)
    print(f"\n✅ 完成！")
    print(f"   输出: {OUTPUT_EPUB}")
    print(f"   大小: {file_size / 1024 / 1024:.1f} MB")
    print(f"   章节: {len(epub_chapters)} | 图片: {img_counter}")

if __name__ == '__main__':
    convert()
